import os

import pandas as pd
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from io import StringIO

from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from langdetect import detect
import fitz
import PyPDF2
import docx

pd.options.display.width = 0


def getListOfFiles(dirName):
    """

    :param dirName: str
        which directory to crawl
    :return: allFiles: List
    list of all the files in directory tree
    """
    # create a list of file and sub directories
    # names in the given directory
    allFiles = list()
    if os.path.basename(dirName) == ".DS_Store":
        return allFiles

    listOfFile = os.listdir(dirName)

    # Iterate over all the entries
    for entry in listOfFile:
        # Create full path
        fullPath = os.path.join(dirName, entry)
        # If entry is a directory then get the list of files in this directory
        if os.path.isdir(fullPath):
            allFiles = allFiles + getListOfFiles(fullPath)
        else:
            allFiles.append(fullPath)

    return allFiles


def check_scan(file_name):
    """
    Calculate the percentage of document that is covered by (searchable) text.

    If the returned percentage of text is very low, the document is
    most likely a scanned PDF
    """

    doc = fitz.open(file_name)
    scanned_pages_ls = []
    scanned_page = False

    pdfFile = PyPDF2.PdfFileReader(file_name)
    # check if file is encrypted
    if pdfFile.isEncrypted:
        try:
            pdfFile.decrypt("")
            print("File Decrypted (PyPDF2)")
        except Exception as ex:
            print("File encrypted could not decrypt")
            scanned_pages_ls.append(f"Encryption error {ex}")
            return scanned_pages_ls

    # check every page if it is a scan
    for page_num, page in enumerate(doc):
        try:
            page_data = pdfFile.getPage(page_num)
        except Exception as ex:
            print(f"Could not find page data {ex}")
            continue
        total_page_area = abs(page.rect)
        text_area = 0.0
        for b in page.get_text_blocks():
            r = fitz.Rect(b[:4])  # rectangle where block text appears
            text_area = text_area + abs(r)
        total_text_area = text_area / total_page_area

        # scanned pages often do not have /Font attribute, not 100% accurate though
        if total_text_area < 0.01 or "/Font" not in page_data["/Resources"]:
            scanned_pages_ls.append(page_num)
            scanned_page = True
            print(f"{page_num} is scanned page")

    doc.close()
    if not scanned_page:
        print("No scanned pages in pdf")
        return "nvt"
    else:
        return len(scanned_pages_ls)


def contract_crawler(path):
    """

    :param path: str
        path to crawl
    :return: data: dictionary
        dictionary with meta_data

    """

    filepaths = [os.path.join(folder, f) for f in os.listdir(path)]

    extensions = []
    file_name_ls = []
    contract_ls = []
    pages = []
    scanned_pages = []
    language = []
    number_files = []
    file_path_ls = []

    for contract in filepaths:

        contract_files = getListOfFiles(contract)

        contract_clean = os.path.basename(contract)

        # if folder is empty
        if not contract_files:
            print(f"Warning: Contract folder for {contract_clean} empty")
            contract_ls.append(contract_clean)
            file_name_ls.append("")
            extensions.append("")
            pages.append("")
            scanned_pages.append("")
            number_files.append("")
            language.append("")
            file_path_ls.append("")
            continue

        print("Number of files in contract dossier: ", len(contract_files))
        print(contract)
        for file in contract_files:
            print("File: ", file)
            file_path_ls.append(file)
            file_name, file_extension = os.path.splitext(file)
            extensions.append(file_extension)

            file_name_ls.append(os.path.basename(file))
            contract_ls.append(contract_clean)
            number_files.append(len(contract_files))

            output_string = StringIO()

            # if document is pdf file
            if file_extension == ".pdf":

                file1 = f"{file}"
                file2 = open(f"{file}", "rb")

                scanned_pages.append(check_scan(file1))
                parser = PDFParser(file2)

                try:
                    document = PDFDocument(parser)
                    rsrcmgr = PDFResourceManager()
                    device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
                    interpreter = PDFPageInterpreter(rsrcmgr, device)

                    for num, page in enumerate(PDFPage.create_pages(document)):
                        if num > 10:
                            break
                        interpreter.process_page(page)
                    pdf_reader = PyPDF2.PdfFileReader(file1)
                    print(pdf_reader.getNumPages())
                    pages.append(pdf_reader.getNumPages())

                    # try to get language of total string of first 10 pages
                    try:
                        language.append(detect(output_string.getvalue()))
                    except Exception as ex:
                        language.append(f"error : {ex}")

                except Exception as ex:
                    print(f"error {ex}")
                    pages.append([ex])
                    language.append([ex])

            # if document is docx file
            elif file_extension == ".docx":

                file1 = f"{file}"
                try:
                    document = docx.Document(file1)
                    lang_str = []
                    for i in document.paragraphs:
                        lang_str.append(i.text)
                    language_str = " ".join(lang_str)
                    # Get language of entire extracted text
                    try:
                        language.append(detect(language_str))
                    except Exception as ex:
                        language.append(f"{ex}")
                    scanned_pages.append("nvt")
                    pages.append("nvt")
                except Exception as ex:
                    scanned_pages.append("nvt")
                    pages.append(f"error {ex}")
                    language.append("nvt")

            else:
                pages.append("nvt")
                language.append("nvt")
                scanned_pages.append("nvt")

        extensions_unique = set(extensions)

        for ex in extensions_unique:
            file_count = extensions.count(ex)
            print(
                f"Found {file_count} files with {ex} extension for contract {contract}"
            )

    data = {
        "Contract": contract_ls,
        "Number of Documents": number_files,
        "File name": file_name_ls,
        "File path": file_path_ls,
        "Type": extensions,
        "Pages": pages,
        "Language": language,
        "Scanned pages": scanned_pages,
    }

    return data


if __name__ == "__main__":
    cur_path = os.path.dirname(__file__)
    folder = os.path.join(cur_path, "..", "../data/raw_data")

    data = contract_crawler(folder)
    df = pd.DataFrame(data)
    df.to_pickle(os.path.join(cur_path, "..", "../data/meta_data_df"))
    df.to_excel(os.path.join(cur_path, "..", "../data/meta_data_volksbank2.xlsx"))
